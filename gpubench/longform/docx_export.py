"""HTML to DOCX, without a Word install and without a third-party converter.

Moved verbatim into the tool. Nothing here knows what is being measured: it walks rendered
HTML and emits Office Open XML. Kept because a reviewer who wants to comment inline needs a
document, not a PDF, and asking them to install a converter loses the review.
"""
import datetime
import os
import re
import sys
import zipfile
from html import unescape
from html.parser import HTMLParser

try:
    import cairosvg
except ImportError:
    cairosvg = None

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Pure text handling, kept in a module with no third-party import so it stays testable
# without the .docx writer present. Re-exported here because callers import it from this module.
from .doctitle import H1_TAG, TITLE_TAG, document_title  # noqa: F401

INK = RGBColor(0x0B, 0x0B, 0x0B)
INK2 = RGBColor(0x3A, 0x39, 0x37)
MUTED = RGBColor(0x6B, 0x69, 0x64)
ACCENT = RGBColor(0x1F, 0x5F, 0xAD)
DRAFT_RED = RGBColor(0x7F, 0x1D, 0x1D)   # the draft stamp, matching the HTML banner
RULE = "D9D8D2"
SURFACE = "F4F3F0"


# --------------------------------------------------------------------------- html model

class Node(object):
    def __init__(self, tag, attrs=None):
        self.tag = tag
        self.attrs = dict(attrs or {})
        self.children = []
        self.text = ""

    def cls(self):
        return self.attrs.get("class", "")


class Parser(HTMLParser):
    """Builds a tree, skipping <svg> subtrees (figures come from the standalone files)."""
    VOID = {"br", "hr", "img", "meta", "link", "input"}

    def __init__(self):
        HTMLParser.__init__(self, convert_charrefs=True)
        self.root = Node("root")
        self.stack = [self.root]
        self.skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if self.skip_depth or tag in ("svg", "style", "script"):
            self.skip_depth += 1
            return
        node = Node(tag, attrs)
        self.stack[-1].children.append(node)
        if tag not in self.VOID:
            self.stack.append(node)

    def handle_startendtag(self, tag, attrs):
        if self.skip_depth:
            return
        self.stack[-1].children.append(Node(tag, attrs))

    def handle_endtag(self, tag):
        if self.skip_depth:
            self.skip_depth -= 1
            return
        if tag in self.VOID:
            return
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]
                return

    def handle_data(self, data):
        if self.skip_depth or not data.strip():
            return
        node = Node("#text")
        node.text = data
        self.stack[-1].children.append(node)


def find(node, tag, cls=None):
    out = []
    for c in node.children:
        if c.tag == tag and (cls is None or cls in c.cls()):
            out.append(c)
        out.extend(find(c, tag, cls))
    return out


def plain(node):
    if node.tag == "#text":
        return node.text
    return "".join(plain(c) for c in node.children)


# --------------------------------------------------------------------------- runs

def add_runs(par, node, bold=False, italic=False, mono=False):
    """Walk inline markup, preserving bold / italic / code / links as Word runs."""
    for c in node.children:
        if c.tag == "#text":
            txt = re.sub(r"\s+", " ", c.text)
            if not txt:
                continue
            run = par.add_run(txt)
            run.bold = bold
            run.italic = italic
            if mono:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            else:
                run.font.color.rgb = INK if bold else INK2
        elif c.tag in ("b", "strong"):
            add_runs(par, c, True, italic, mono)
        elif c.tag in ("i", "em"):
            add_runs(par, c, bold, True, mono)
        elif c.tag == "code":
            add_runs(par, c, bold, italic, True)
        elif c.tag == "a":
            before = len(par.runs)
            add_runs(par, c, bold, italic, mono)
            for run in par.runs[before:]:
                run.font.color.rgb = ACCENT
                run.underline = True
            href = c.attrs.get("href", "")
            if href.startswith("http"):
                par.add_run(" (%s)" % href).font.color.rgb = MUTED
        elif c.tag == "br":
            par.add_run().add_break()
        else:
            add_runs(par, c, bold, italic, mono)


def shade(cell, hexcolor):
    tc = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    tc.append(el)


def borderless_box(table_obj, color=RULE):
    tbl = table_obj._tbl
    pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    pr.append(borders)


# --------------------------------------------------------------------------- blocks

def _fit_to_page(table_obj):
    """Constrain a table to 100% of the text column and turn OFF fixed layout.

    tblLayout=autofit lets Word size columns from content and wrap within them, rather than
    honouring widths that may exceed the page. Belt and braces with autofit=True, because the two
    settings are honoured by different Word versions.
    """
    tbl = table_obj._tbl
    pr = tbl.tblPr
    for tag, attrs in (("w:tblW", {"w:w": "5000", "w:type": "pct"}),
                       ("w:tblLayout", {"w:type": "autofit"})):
        for existing in pr.findall(qn(tag)):
            pr.remove(existing)
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        pr.append(el)


def repeat_header(row):
    """Make a header row repeat on every page the table spans.

    Bold and shaded makes a row LOOK like a header on the page it starts on. Word only REPEATS it
    across a page break when the row's properties carry w:tblHeader, and without that a reader who
    scrolls into the continuation of a long table sees unlabelled columns. This report has several
    tables longer than a page, so it is not hypothetical.

    Written as XML because python-docx 1.2.0 has no such property: assigning
    `row.repeat_as_header_row = True` silently creates a new Python attribute and emits nothing,
    which is how the first attempt at this "passed" while changing zero of 64 tables. Found by
    unzipping the shipped .docx and counting w:tblHeader, not by reading the build log.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def render_table(doc, node):
    heads = [plain(th).strip() for th in find(node, "th")]
    rows = []
    for tr in find(node, "tr"):
        tds = [td for td in tr.children if td.tag == "td"]
        if tds:
            rows.append(tds)
    if not heads and not rows:
        return
    ncols = max(len(heads), max((len(r) for r in rows), default=0))
    t = doc.add_table(rows=1 if heads else 0, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Fit the table to the text column and let cells WRAP. Word wraps by default, but an
    # unconstrained table can still run past the right margin on a wide row; pinning the width to
    # the page and enabling autofit keeps every column on the page. A reader should never have to
    # discover that a column was cut off.
    t.autofit = True
    _fit_to_page(t)
    borderless_box(t)
    if heads:
        for i, htext in enumerate(heads[:ncols]):
            cell = t.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(htext)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = INK
            shade(cell, SURFACE)
        repeat_header(t.rows[0])
    for tds in rows:
        cells = t.add_row().cells
        for i, td in enumerate(tds[:ncols]):
            cell = cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, td)
            for r in p.runs:
                r.font.size = Pt(9)
    doc.add_paragraph()


def render_figure(doc, node, figdir, pngdir, counter):
    fid = node.attrs.get("id", "")
    svg = os.path.join(figdir, fid + ".svg")
    png = os.path.join(pngdir, fid + ".png")
    if os.path.exists(svg) and cairosvg is not None:
        if not os.path.exists(png) or os.path.getmtime(png) < os.path.getmtime(svg):
            cairosvg.svg2png(url=svg, write_to=png, scale=2.0, background_color="#ffffff")
    if os.path.exists(png):
        doc.add_picture(png, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cap in find(node, "figcaption"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        add_runs(p, cap)
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = MUTED
        break
    for note in find(node, "p", "note"):
        p = doc.add_paragraph()
        add_runs(p, note)
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = MUTED
        break
    # the table view that accompanies every figure
    for t in find(node, "table"):
        render_table(doc, t)
        break


def render(doc, node, figdir, pngdir, counter):
    for c in node.children:
        tag, cls = c.tag, c.cls()
        if tag == "h1":
            h = doc.add_heading(plain(c).strip(), level=0)
            for r in h.runs:
                r.font.color.rgb = INK
        elif tag == "h2":
            doc.add_page_break()
            h = doc.add_heading(plain(c).strip(), level=1)
            for r in h.runs:
                r.font.color.rgb = INK
        elif tag == "h3":
            h = doc.add_heading(plain(c).strip(), level=2)
            for r in h.runs:
                r.font.color.rgb = INK
        elif tag == "figure":
            counter[0] += 1
            render_figure(doc, c, figdir, pngdir, counter)
        elif tag == "table":
            render_table(doc, c)
        elif tag == "pre":
            p = doc.add_paragraph()
            run = p.add_run(plain(c).rstrip())
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Inches(0.25)
        elif tag == "ul" or tag == "ol":
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in [x for x in c.children if x.tag == "li"]:
                p = doc.add_paragraph(style=style)
                add_runs(p, li)
        elif tag == "dl":
            # A DEFINITION LIST HAD NO BRANCH AT ALL, so it fell through to the recursion below,
            # which walks into <dt> and <dd>, finds no branch for those either, walks into their
            # #text children, and drops every one: the block vanished from the Word edition
            # entirely while the HTML kept it. It was found on a title-page control block whose
            # two numbers were the only figures that document computed about itself, and worked
            # around in that document rather than fixed here, which left it live for every other
            # document this module converts.
            for item in c.children:
                if item.tag == "dt":
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    add_runs(p, item, bold=True)
                elif item.tag == "dd":
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    add_runs(p, item)
        elif tag == "div" and "draftbanner" in cls:
            # The draft stamp. A DOCX is the format that gets mailed and commented on, so a
            # document the gate did not pass has to say so in Word too, at the top, in red.
            for sub in find(c, "p"):
                p = doc.add_paragraph()
                add_runs(p, sub)
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(11)
                    r.font.color.rgb = DRAFT_RED
                break
        elif tag == "div" and "stats" in cls:
            # stat tiles become a compact two-column table
            pairs = []
            for s in find(c, "div", "stat"):
                n = find(s, "div", "n")
                k = find(s, "div", "k")
                if n and k:
                    pairs.append((plain(n[0]).strip(), plain(k[0]).strip()))
            if pairs:
                t = doc.add_table(rows=0, cols=2)
                # Same page-fitting as every other table. This one is built by a separate path,
                # which is exactly how it got missed the first time: 51 tables constrained, 1 not.
                t.autofit = True
                _fit_to_page(t)
                borderless_box(t)
                for n, k in pairs:
                    cells = t.add_row().cells
                    r = cells[0].paragraphs[0].add_run(n)
                    r.bold = True
                    r.font.size = Pt(13)
                    r.font.color.rgb = INK
                    r2 = cells[1].paragraphs[0].add_run(k)
                    r2.font.size = Pt(9)
                    r2.font.color.rgb = MUTED
                doc.add_paragraph()
        elif tag == "div" and "callout" in cls:
            for sub in c.children:
                if sub.tag == "p":
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    add_runs(p, sub)
                    shade_paragraph(p)
        elif tag == "p":
            p = doc.add_paragraph()
            add_runs(p, c)
            if "lede" in cls:
                for r in p.runs:
                    r.font.size = Pt(12)
                    r.font.color.rgb = INK
            elif "note" in cls or "meta" in cls:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = MUTED
        elif tag == "hr":
            doc.add_paragraph("_" * 60).runs[0].font.color.rgb = RGBColor(0xD0, 0xCF, 0xC9)
        elif tag in ("details", "summary"):
            render(doc, c, figdir, pngdir, counter)
        else:
            render(doc, c, figdir, pngdir, counter)


def shade_paragraph(p):
    pr = p._p.get_or_add_pPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), SURFACE)
    pr.append(el)


def main(argv=None, title=None, subject=""):
    """Convert a rendered HTML report to DOCX.

    Paths are derived from the SOURCE DOCUMENT, not from this file. That distinction matters now
    that this module lives in the tool rather than beside one report: deriving `here` from
    __file__ would send it hunting for figures inside the tool's own package directory.

    `title` and `subject` go into the Word core properties. The title defaults to the DOCUMENT'S
    OWN title rather than to a constant, and the subject to nothing at all: a converter that
    cannot read a subject off the page has no business inventing one, and an empty field is the
    honest answer where a borrowed sentence is a false one.
    """
    argv = list(sys.argv[1:] if argv is None else argv)
    src = argv[0] if argv else None
    if not src:
        raise SystemExit("usage: docx_export <rendered-report.html>")
    src = os.path.abspath(src)
    here = os.path.dirname(src)
    figdir = os.path.join(here, "figures")
    pngdir = os.path.join(here, "figures", "png")
    if not os.path.isdir(pngdir):
        os.makedirs(pngdir)
    if cairosvg is None:
        print("warning: cairosvg missing, figures will be omitted", file=sys.stderr)

    with open(src, "r", encoding="utf-8") as f:
        html = f.read()
    parser = Parser()
    parser.feed(html)

    body = parser.root
    mains = find(body, "main")
    if mains:
        body = mains[0]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(8)
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    render(doc, body, figdir, pngdir, [0])

    # An exported document should not advertise the toolchain that produced it, and should not
    # carry claims that are simply untrue.
    cp = doc.core_properties
    cp.author = "Muhammad Asif"
    cp.last_modified_by = "Muhammad Asif"
    cp.comments = ""
    cp.category = ""
    cp.keywords = ""
    cp.title = title if title is not None else document_title(
        html, fallback=os.path.splitext(os.path.basename(src))[0])
    cp.subject = subject or ""
    now = datetime.datetime.now()
    cp.created = now
    cp.modified = now
    cp.revision = 1

    out = os.path.splitext(src)[0] + ".docx"
    doc.save(out)
    scrub_package(out)
    print("wrote %s (%.0f KB)" % (out, os.path.getsize(out) / 1024.0))
    return 0


def scrub_package(path):
    """Clean docProps/app.xml, which python-docx does not expose.

    The default template ships `<Application>Microsoft Macintosh Word</Application>` and a 2013
    template date. Neither is true of this file, and shipping a document that misstates how it was
    produced is worse than shipping one that says nothing.
    """
    tmp = path + ".tmp"
    with zipfile.ZipFile(path, "r") as zin:
        items = [(i, zin.read(i.filename)) for i in zin.infolist()]
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for info, data in items:
            if info.filename == "docProps/app.xml":
                text = data.decode("utf-8", "replace")
                text = re.sub(r"<Application>.*?</Application>", "<Application></Application>", text)
                text = re.sub(r"<Template>.*?</Template>", "<Template></Template>", text)
                text = re.sub(r"<AppVersion>.*?</AppVersion>", "", text)
                text = re.sub(r"<Manager/>|<Manager>.*?</Manager>", "<Manager/>", text)
                text = re.sub(r"<Company/>|<Company>.*?</Company>", "<Company/>", text)
                data = text.encode("utf-8")
            zout.writestr(info, data)
    os.replace(tmp, path)


if __name__ == "__main__":
    sys.exit(main())
